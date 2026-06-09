
## 在绘图的命令窗格中执行
import os
import time
from pathlib import Path
import labrad
import matplotlib
import numpy as np
matplotlib.rcParams['backend'] = 'Agg'
import matplotlib.pyplot as plt
from io import BytesIO  # 用于创建内存缓冲区
from PIL import Image   # 用于读取和处理图片
from lqms.pyle import registry_wrapper2

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False
import io
from docx import Document
from docx.shared import Inches, Pt
from lqms.data_process import data_process as dp
from lqms.data_process.dataAnalysisCore import DataLab
from lqms.data_process.xeb import process_xeb as px
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lqms.data_process import *


import os
import pathlib
import re
import sys
import time
from importlib import reload

import labrad

# basic
import ray
from labrad.units import GHz, MHz, V, dBm, kHz, mV, ns, rad, us
from lqcs import system_config
from pylab import np, plt

from lqms.data_process import dataAnalysisCore as dc
from lqms.measure import (
    generate_coupler,
    generate_qubit,
)
from lqms.measure.basic import BasicTuner, util
from lqms.measure.tuners import base_experiment as basex
from lqms.measure.tuners import cz_experiment as czx
from lqms.measure.tuners import cz_nodes as czn
from lqms.measure.tuners import smt_experiment
from lqms.measure.tuners import sq_nodes as sq
from lqms.pyle.util import sweeptools as st
from lqms.pyle.workflow import switchSession
from lqms.utils import rect_lattice as rl
from lqms.utils.save_path import get_info_path





import win32gui
import win32ui
import win32con
import win32api
from ctypes import windll

def capture_screen_precise(filename):
    """
    使用更精确的方法获取屏幕尺寸
    """
    try:
        # 获取桌面窗口
        hdesktop = win32gui.GetDesktopWindow()
        
        # 方法1：获取窗口的矩形区域
        rect = win32gui.GetWindowRect(hdesktop)
        window_width = rect[2] - rect[0]
        window_height = rect[3] - rect[1]
        
        # 方法2：获取系统指标
        system_width = win32api.GetSystemMetrics(win32con.SM_CXSCREEN)
        system_height = win32api.GetSystemMetrics(win32con.SM_CYSCREEN)
        
        # 方法3：考虑 DPI 缩放
        user32 = windll.user32
        user32.SetProcessDPIAware()
        dpi_width = user32.GetSystemMetrics(win32con.SM_CXSCREEN)
        dpi_height = user32.GetSystemMetrics(win32con.SM_CYSCREEN)
        
        print(f"窗口尺寸: {window_width}x{window_height}")
        print(f"系统尺寸: {system_width}x{system_height}")
        print(f"DPI 尺寸: {dpi_width}x{dpi_height}")
        
        # 使用最大的尺寸
        width = max(window_width, system_width, dpi_width)
        height = max(window_height, system_height, dpi_height)
        
        print(f"最终使用尺寸: {width}x{height}")
        
        # 创建设备上下文
        desktop_dc = win32gui.GetWindowDC(hdesktop)
        img_dc = win32ui.CreateDCFromHandle(desktop_dc)
        mem_dc = img_dc.CreateCompatibleDC()
        
        # 创建位图对象
        screenshot = win32ui.CreateBitmap()
        screenshot.CreateCompatibleBitmap(img_dc, width, height)
        mem_dc.SelectObject(screenshot)
        
        # 复制屏幕内容到位图
        mem_dc.BitBlt((0, 0), (width, height), img_dc, (0, 0), win32con.SRCCOPY)
        
        # 保存为 BMP 文件
        screenshot.SaveBitmapFile(mem_dc, filename)
        
        print(f"截图已保存: {filename}")
        
    except Exception as e:
        print(f"截图失败: {e}")
        
    finally:
        # 清理资源
        try:
            mem_dc.DeleteDC()
            win32gui.DeleteObject(screenshot.GetHandle())
            win32gui.ReleaseDC(hdesktop, desktop_dc)
        except:
            pass




def create_graphs(data_orders, exp_name):
    charts = []
    qobjs = []
    for order in data_orders:
        # qubit = re.search(pattern, file_names[order-1]).group()
        # print(order)
        data.loadDataset(int(order))
        timestamp = data.parameters['create_time']
        qobj = data.parameters['config'][0]
        name = dv.dir()[1][order - 1]
        qobjs.append(qobj)
        # print(order-1, len(dv.dir()[1]))
        fread = np.round(data.parameters[f'{qobj}.fread'], 4)
        f10 = np.round(data.parameters[f'{qobj}.f10'], 4)
        bias = np.round(data.parameters[f'{qobj}.bias_z'], 2)
        if 'XEB' in data.dataset_name:
            xeb_res = px.XEB(data, [int(order)], collect=True)
            xeb_fid = 1-xeb_res[int(order)]['error_Pauli_per_cycle']/1.5
            fids[qobj] = xeb_fid
        elif 'T1' in data.dataset_name:
            t1_res = dp.T1(data)
            T1s[qobj] = t1_res[1][1][0]
            f10s[qobj] = f10
            freads[qobj] = fread
        elif 'S21_dis' in data.dataset_name:
            dp.dispersive_shift(data)
        elif 'PiPulse' in data.dataset_name:
            data.plotDataset(yIdx=5)
        elif 'S21power2d' in data.dataset_name:
            data.plotDataset(yIdx=1)
        elif 'Spectroscopy' in data.dataset_name:
            data.plotDataset(yIdx=1)
        elif 'IQraw' in data.dataset_name:
            continue
        else:
            data.plotDataset()
        plt.annotate(
            f'Obtaining time: {timestamp}',  # 文本内容
            xy=(0.95, 0.05),  # 注释位置（右下角，x=0.95 靠右，y=0.05 靠下）
            fontsize=20,
            xycoords='axes fraction',  # 坐标以图像比例为单位（0-1）
            bbox=dict(
                boxstyle='round,pad=0.3',  # 圆角矩形，内边距0.3
                facecolor='white',  # 背景色白色
                edgecolor='none',  # 无边框
                alpha=0.8,  # 半透明（避免完全遮挡背景图像）
            ),
            ha='right',  # 水平对齐方式：右对齐
            va='bottom',  # 垂直对齐方式：底部对齐
        )
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        if 'Spectroscopy' in data.dataset_name:
            charts.append((f'{qobj}实验运行结果（{timestamp}），比特频率：{f10}', img_buffer))
        else:
            charts.append((f'{qobj}实验运行结果（{timestamp}）', img_buffer))
        plt.close()

    return charts

def png_to_img_buffer(png_path):
    """
    将 PNG 文件转换为内存中的 BytesIO 缓冲区（img_buffer）
    :param png_path: PNG 文件的路径（相对或绝对路径）
    :return: 包含图片数据的 BytesIO 对象
    """
    # 1. 打开 PNG 图片
    with Image.open(png_path) as img:
        # 2. 创建内存缓冲区
        img_buffer = BytesIO()
        # 3. 将图片数据写入缓冲区（格式指定为 PNG）
        img.save(img_buffer, format='PNG')
        # 4. 将缓冲区指针重置到起始位置（否则读取时可能为空）
        img_buffer.seek(0)
    return img_buffer

def create_docx_with_charts(charts, filename='zpa2d.docx'):
    """创建包含图表的Word文档"""
    # 创建新文档
    try:
        # 1. 打开现有文档（关键：不是新建，而是读取已有文档）
        doc = Document(filename)
    except:
        # 若文件不存在，创建新文档
        doc = Document()
        print(f"文件 {filename} 不存在，将创建新文档并添加内容")

    # 添加标题
    # title = doc.add_heading('测试报告', 0)
    # title.alignment = 1  # 居中对齐
    
    # 添加每个图表到文档中
    for chart_name, img_buffer in charts:
        # 添加图表标题
        doc.add_heading(f'量子比特编号: {chart_name}', level=1)
        # 插入图片
        if("截图" in chart_name):
            picture_paragraph  = doc.add_picture(img_buffer, width=Inches(6))
        else:
            picture_paragraph  = doc.add_picture(img_buffer, width=Inches(4))
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 添加空行分隔
        # doc.add_paragraph()

    # 保存文档
    doc.save(filename)
    print(f'文档已保存为: {filename}')


def plot_data(datas, names):
    charts = []
    for i,data in enumerate(datas):
        # 排序数据（按数值从小到大，使条形图更易对比）
        sorted_items = sorted(data.items(), key=lambda x: x[1])
        labels, values = zip(*sorted_items)  # 拆分标签和数值

        # 设置画布
        plt.figure(figsize=(20, 10))  # 宽20，高10，适配较多标签

        # 绘制横向条形图
        bars = plt.bar(labels, values, color='skyblue', edgecolor='black')

        # 添加数值标签（在条形右侧显示具体值）
        for bar in bars:
            height = bar.get_height()
            plt.text(
            bar.get_x() + bar.get_width()/2,  # 水平居中
            height,  # 条形顶部+3偏移，避免重叠
            f'{height:.4f}',  # 保留1位小数
            ha='center',  # 水平对齐
            fontsize=20,
            rotation=30
            )

        # 设置标题和坐标轴标签
        plt.title(f'各量子比特{names[i]}对比', fontsize=50, pad=20)
        plt.xlabel(f'{names[i]}', fontsize=45, labelpad=10)
        plt.ylabel('量子比特标签', fontsize=45, labelpad=10)

        # 调整刻度字体大小
        plt.xticks(rotation=45, ha='right', fontsize=20)
        plt.yticks(fontsize=45)

        # 去除上.右边框，使图表更简洁
        plt.gca().spines['top'].set_visible(False)
        plt.gca().spines['right'].set_visible(False)

        # 调整布局，避免标签被截断
        plt.tight_layout()
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        charts.append((f' ', img_buffer))
    return charts

def gen_overview_docx(datas, names, filename):
    # 创建Word文档
    # 创建新文档
    try:
        # 1. 打开现有文档（关键：不是新建，而是读取已有文档）
        doc = Document(filename)
    except:
        # 若文件不存在，创建新文档
        doc = Document()
        print(f"文件 {filename} 不存在，将创建新文档并添加内容")
    # 添加标题
    title = doc.add_heading('量子比特测试结果汇总', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中
    charts = plot_data(datas, names)
    data = datas[-2]
    # 数据排序（按保真度升序）
    sorted_data = sorted(zip(data.keys(), data.values()), key=lambda x: x[1])
    # 添加表格（行数=数据行数+1表头，列数=2）
    table = doc.add_table(rows=1, cols=5, style='Table Grid')
    table.autofit = False  # 关闭自动适应，手动设置列宽
    table.columns[0].width = Inches(1.5)  # 第一列（标签）宽度
    table.columns[1].width = Inches(1)    # 第二列（参数值）宽度
    table.columns[2].width = Inches(1)    # 第二列（参数值）宽度
    table.columns[3].width = Inches(1)    # 第二列（参数值）宽度
    table.columns[4].width = Inches(1)    # 第二列（参数值）宽度
    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '量子比特标签'
    hdr_cells[1].text = names[0]
    hdr_cells[2].text = names[1]
    hdr_cells[3].text = names[2]
    hdr_cells[4].text = names[3]
    # 格式化表头文字（加粗.居中）
    for cell in hdr_cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.bold = True
            run.font.size = Pt(10)
    # 填充表格数据
    for label, value in sorted_data:
        row_cells = table.add_row().cells
        # 第一列：标签
        row_cells[0].text = label
        # 第二列：参数值（保留2位小数）
        try:
            row_cells[1].text = f"{datas[0][label]:.4f}"
        except:
            pass
        try:
            row_cells[2].text = f"{datas[1][label]:.4f}"
        except:
            pass
        try:
            row_cells[3].text = f"{datas[2][label]:.4f}"
        except:
            pass
        try:
            row_cells[4].text = f"{datas[3][label]:.4f}"
        except:
            pass
        
        # 格式化单元格文字（居中，调整字体大小）
        for cell in row_cells:
            for para in cell.paragraphs:
                if not para.runs:  # 判断 runs 列表是否为空
                    para.add_run()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0]
                run.font.size = Pt(9)

     # 添加每个图表到文档中
    for chart_name, img_buffer in charts:
        # 添加图表标题
        doc.add_heading(f'{chart_name}', level=1)

        # 插入图片
        if("截图" in chart_name):
            picture_paragraph  = doc.add_picture(img_buffer, width=Inches(8.8))
        else:
            picture_paragraph  = doc.add_picture(img_buffer, width=Inches(6))
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 添加空行分隔
        # doc.add_paragraph()


    # 保存文档
    doc.save(filename)
    print(f'文档已保存为: {filename}')



freads = {}
f10s = {}
T1s = {}
fids = {}
cxn = labrad.connect()
util.setWiringInfo(cxn)
user = 'LQHL'
s = switchSession(cxn, user=user)
session = ['', 'LQHL', 'test', '20260324']
dv = cxn.data_vault

data = dc.DataLab(session, dv, dv_type='data_vault')
# info_path = get_info_path(s)
# info = dc.InfoBase(info_path)
try:
    info = dc.InfoBase(get_info_path(s))
except Exception:
    info = None
qter = QubitUpdater(data, info)


# experiment
try:
    info = dc.InfoBase(get_info_path(s))
except Exception:
    info = None
_all_qubits = generate_qubit(globals(), info=info, sample=s)
_all_couplers = generate_coupler(globals(), info=info, sample=s)
# for impa in DEVICE.IMPAS.values():
#     exec(f'{impa} = qubitAuto.Qubit("{impa}")')

auto_config = {
    'stats': 300,
    'correctX': False,
    'correctZ': False,
    'reset': False,
    'apply_21': False,
    'run_mode': 'local',
}
# initialize basic tunner
BasicTuner(**auto_config)
BasicTuner._sample = s
BasicTuner._all_qobjs = _all_qubits | _all_couplers


# ray
HEAD_IP_ADDRESS = system_config.get_ray_head()
NODE_IP_ADDRESS = system_config.get_config()['ip']
RAY_PORT = system_config.get_ray_port()
print('NODE_IP_ADDRESS', NODE_IP_ADDRESS)

if not ray.is_initialized():
    ray.init(
        address=f'{HEAD_IP_ADDRESS}:{RAY_PORT}',
        namespace='main',
        _node_ip_address=NODE_IP_ADDRESS,
        log_to_driver=False,
    )
def write_cfg(order):
    dv = cxn.data_vault
    session = ['', 'LQHL', 'test', '20260324']
    dv.cd(session)
    data = dc.DataLab(session, dv, dv_type='data_vault')
    data.loadDataset(int(order))
    timestamp = data.parameters['create_time']
    qobj = data.parameters['config'][0]
    name = dv.dir()[1][order - 1]
    fread = np.round(data.parameters[f'{qobj}.fread'], 5)
    f10 = np.round(data.parameters[f'{qobj}.f10'], 5)
    bias = np.round(data.parameters[f'{qobj}.bias_z'], 3)
    read_power = np.round(data.parameters[f'{qobj}.ReadIn.power'], 3)
    pi_length = np.round(data.parameters[f'{qobj}.PiGate.length'], 3)
    pi_amp = np.round(data.parameters[f'{qobj}.PiGate.amp'], 3)
    

    s[qobj].fread = data.parameters[f'{qobj}.fread']
    s[qobj].f10 = data.parameters[f'{qobj}.f10']
    s[qobj].fc = data.parameters[f'{qobj}.fc']
    s[qobj].f21 = data.parameters[f'{qobj}.f21']
    s[qobj].bias_z = data.parameters[f'{qobj}.bias_z']


    s[qobj].PiGate.alpha = data.parameters[f'{qobj}.PiGate.alpha']
    s[qobj].PiGate.amp = data.parameters[f'{qobj}.PiGate.amp']
    s[qobj].PiGate.length = data.parameters[f'{qobj}.PiGate.length']
    s[qobj].PiGate.zpa = data.parameters[f'{qobj}.PiGate.zpa']

    s[qobj].PiHalf.amp = data.parameters[f'{qobj}.PiHalf.amp']
    s[qobj].PiHalf.alpha = data.parameters[f'{qobj}.PiHalf.alpha']
    s[qobj].PiHalf.length = data.parameters[f'{qobj}.PiHalf.length']
    s[qobj].PiHalf.zpa = data.parameters[f'{qobj}.PiHalf.zpa']



    s[qobj].ReadIn.length = data.parameters[f'{qobj}.ReadIn.length']
    s[qobj].ReadIn.power = data.parameters[f'{qobj}.ReadIn.power']
    s[qobj].ReadIn.ring_power = data.parameters[f'{qobj}.ReadIn.ring_power']
    s[qobj].ReadIn.ring_length = data.parameters[f'{qobj}.ReadIn.ring_length']
    s[qobj].ReadIn.zpa = data.parameters[f'{qobj}.ReadIn.zpa']

    s[qobj].ReadOut.amp = data.parameters[f'{qobj}.ReadOut.amp']
    s[qobj].ReadOut.length = data.parameters[f'{qobj}.ReadOut.length']
    s[qobj].ReadOut.window_type = data.parameters[f'{qobj}.ReadOut.window_type']

    s[qobj].PiGate.length = data.parameters[f'{qobj}.PiGate.length']
    s[qobj].PiGate.amp = data.parameters[f'{qobj}.PiGate.amp']

    s[qobj].discriminator.center0 = data.parameters[f'{qobj}.discriminator.center0']
    s[qobj].discriminator.center1 = data.parameters[f'{qobj}.discriminator.center1']
    s[qobj].discriminator.measure_f0 = data.parameters[f'{qobj}.discriminator.measure_f0']
    s[qobj].discriminator.measure_f1 = data.parameters[f'{qobj}.discriminator.measure_f1']
    s[qobj].discriminator.method = data.parameters[f'{qobj}.discriminator.method']
    s[qobj].discriminator.radius0 = data.parameters[f'{qobj}.discriminator.radius0']
    s[qobj].discriminator.threshold = data.parameters[f'{qobj}.discriminator.threshold']



if __name__ == '__main__':
    name = 'fzq'
    time_stamp = time.strftime('%Y%m%d-%H_%M_%S')
    today = time.strftime('%Y%m%d')
    Path(f"./capture_screen/").mkdir(parents=True, exist_ok=True)
    Path(f"./test_{today}_by_{name}").mkdir(parents=True, exist_ok=True)
    capture_screen_filename = f"./capture_screen/capture_screen_{time_stamp}.png"
    capture_screen_precise(capture_screen_filename)
    capture_screen_buffer =  png_to_img_buffer(capture_screen_filename)

    # qobj_tuples = [(q4ru3,41610),(q8ru2,41148),(q10ru2,41034),(q1ru3,40924),(q12ru4,40718),(q10ru3,40490),
    #                (q11ru5,40439),(q6ru5,40412),(q5ru4,40350),(q12ru5,40333),(q5lu5,40216),(q9ru5,37839),
    #                (q7ru4,37808),(q11lu4, 61467),(q12ru5,40333)]

    qobj_tuples = [(q2lu4, -1),(q12lu4,-1)]


    for qobj, formal_order in qobj_tuples:

        # write_cfg(formal_order)


        # ## experiments
        # try:
        #     sq.iqraw(qobj,do_plot=False)
        #     fid = qter.fitData(-1, collect=True, do_plot=False)[1][-2]
        # except:
        #     continue
        # if(fid<0.2):
        #     print("读取区分度过低，请校准后重试")
        #     continue

        ## 指标采集
        sq.piamp(qobj,amp=3, update=False,do_plot=False)
        sq.t1(qobj,zpa=0,do_plot=False)
        sq.t1(qobj,zpa=qobj.regs.bias_z,do_plot=False)
        sq.spectroscopy(qobj, freq=np.arange(-0.1, 0.1, 0.002)+qobj.regs.f10,spec_amp=1, update=False, do_plot=False)
        sq.s21_dis(qobj,do_plot=False, update=False)
        sq.xeb(qobj)


        ##
        today = time.strftime('%Y%m%d')
        
        start_ds = -7  # 开始的dataset 编号
        
        

        cxn = labrad.connect()
        dv = cxn.data_vault
        data = DataLab(session, dv, dv_type='data_vault') 
        dv.cd(session)
        results = dv.get_tags(([], dv.dir()[1][start_ds:]))[1]
        
        
        for exp_name in ['T1','PiPulse','Spectroscopy','XEB reference','S21_dis']:
            data_orders = np.array([int(da[0][:5]) for da in results if ((exp_name in da[0]))])
            data_orders = data_orders[data_orders >= start_ds]
            # print(dv.dir())
            # # 示例1: 创建包含多个图表的报告
            print('正在创建示例图表...')
            # print(capture_screen_buffer)
            charts = [(f"{str(qobj)}开始运行截图（"+time_stamp+"）", capture_screen_buffer)]
            experiment_charts = create_graphs(data_orders[::-1], exp_name)
            charts.extend(experiment_charts)

            filename = f"./test_{today}_by_{name}/report_of_{exp_name}.docx"
            create_docx_with_charts(charts, filename)
    datas = [freads, f10s, T1s, fids]
    names=['fread(GHz)', 'f10(GHz)','T1(μs)','xeb_fid']
    gen_overview_docx(datas, names, f"./test_{today}_by_{name}/overview_report_by_{name}.docx")
    print('所有操作完成！')
